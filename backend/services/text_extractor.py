"""
services/text_extractor.py

Extracts plain text from any document format the CRPF platform accepts:
  - PDF (text-based, clean)        → pdfplumber (fast, no ML needed)
  - PDF (scanned / image-based)    → docTR OCR fallback
  - DOCX / Word                    → python-docx
  - Images (JPG, PNG, TIFF, etc.)  → docTR OCR

Returns a single clean UTF-8 string ready to be stored and forwarded
to the evaluation stage.

Design decisions:
  - For PDFs we always try pdfplumber first (milliseconds). If it returns
    fewer than MIN_TEXT_CHARS characters the document is likely scanned,
    so we fall back to docTR OCR (seconds).
  - docTR is imported lazily (inside the function) to avoid loading the
    ~300MB model at import time on every worker restart. On first call it
    loads once and stays in memory for the process lifetime.
  - All errors are caught and re-raised as ExtractionError so the router
    can return a clean 422 to the frontend.
  - Every extracted string is stripped of null bytes and excessive whitespace
    so it is safe to store in PostgreSQL TEXT columns.
"""

import io
import logging
import re
from typing import Tuple

logger = logging.getLogger(__name__)

# Minimum characters from pdfplumber before we assume the PDF is scanned
MIN_TEXT_CHARS = 100


class ExtractionError(Exception):
    """Raised when text extraction fails for any reason."""
    pass


# ── Helpers ──────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """
    Remove null bytes, collapse runs of whitespace, strip leading/trailing space.
    PostgreSQL TEXT columns reject null bytes; excessive whitespace wastes storage
    and makes LLM prompts harder to reason about.
    """
    text = text.replace("\x00", "")              # null bytes
    text = re.sub(r"[ \t]+", " ", text)          # collapse horizontal space
    text = re.sub(r"\n{3,}", "\n\n", text)       # collapse blank lines
    return text.strip()


def _ocr_with_doctr(file_bytes: bytes, mime_type: str) -> str:
    """
    Run docTR OCR on raw bytes.
    Supports both PDF and image inputs via docTR's DocumentFile helpers.
    Model is loaded once per process (lazy singleton pattern).
    """
    try:
        from doctr.io import DocumentFile
        from doctr.models import ocr_predictor
    except ImportError:
        raise ExtractionError(
            "docTR is not installed. Add 'python-doctr[torch]' to requirements.txt "
            "and rebuild the Docker image."
        )

    # Lazy-load the OCR model — expensive on first call (~5s), free after that
    if not hasattr(_ocr_with_doctr, "_model"):
        logger.info("Loading docTR OCR model (first call — this takes ~5s)...")
        _ocr_with_doctr._model = ocr_predictor(pretrained=True)
        logger.info("docTR model loaded and cached.")

    model = _ocr_with_doctr._model

    # Build DocumentFile from bytes depending on format
    is_image_mime = mime_type.startswith("image/")
    if is_image_mime:
        doc = DocumentFile.from_images([file_bytes])
    else:
        # Treat as PDF (works for both native PDF and scanned PDF)
        doc = DocumentFile.from_pdf(file_bytes)

    result = model(doc)

    # Flatten the nested docTR output into plain text
    # Structure: result.pages → blocks → lines → words
    lines = []
    for page in result.pages:
        for block in page.blocks:
            for line in block.lines:
                line_text = " ".join(word.value for word in line.words)
                lines.append(line_text)

    return "\n".join(lines)


# ── Per-format extractors ─────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> str:
    """
    Try pdfplumber first (text-based PDF).
    Fall back to docTR if the PDF is scanned / image-only.
    """
    try:
        import pdfplumber
    except ImportError:
        raise ExtractionError(
            "pdfplumber is not installed. Add it to requirements.txt."
        )

    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as exc:
        logger.warning("pdfplumber failed (%s) — will try docTR OCR", exc)
        text_parts = []

    combined = "\n".join(text_parts)

    if len(combined.strip()) >= MIN_TEXT_CHARS:
        logger.info("PDF extracted via pdfplumber (%d chars).", len(combined))
        return combined

    # Not enough text → scanned PDF, fall back to OCR
    logger.info(
        "pdfplumber returned %d chars (< %d) — PDF appears scanned, switching to docTR OCR.",
        len(combined.strip()), MIN_TEXT_CHARS
    )
    ocr_text = _ocr_with_doctr(file_bytes, "application/pdf")
    logger.info("docTR OCR returned %d chars for scanned PDF.", len(ocr_text))
    return ocr_text


def _extract_docx(file_bytes: bytes) -> str:
    """Extract all paragraph text from a DOCX file using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ExtractionError(
            "python-docx is not installed. Add 'python-docx' to requirements.txt."
        )

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ExtractionError(f"Could not open DOCX: {exc}") from exc

    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # Also extract text from tables (common in Indian tender documents)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    return "\n".join(paragraphs)


def _extract_image(file_bytes: bytes, mime_type: str) -> str:
    """Run docTR OCR directly on an image file."""
    return _ocr_with_doctr(file_bytes, mime_type)


# ── Public API ────────────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, mime_type: str, original_filename: str) -> Tuple[str, str]:
    """
    Extract plain text from a document.

    Parameters
    ----------
    file_bytes        : raw bytes of the uploaded file
    mime_type         : MIME type string from the upload (e.g. 'application/pdf')
    original_filename : used as a fallback to infer format when MIME is generic

    Returns
    -------
    (extracted_text, method_used)
      extracted_text : clean UTF-8 string, ready for DB storage and LLM input
      method_used    : one of "pdfplumber", "doctr_ocr", "python_docx", "doctr_image"
                       — stored in the DB for audit purposes

    Raises
    ------
    ExtractionError : if the format is unsupported or extraction fails
    """
    fname_lower = (original_filename or "").lower()

    # Determine format — MIME type is preferred, filename extension is fallback
    # (browsers sometimes send 'application/octet-stream' for everything)
    is_pdf = (
        mime_type == "application/pdf"
        or fname_lower.endswith(".pdf")
    )
    is_docx = (
        mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        )
        or fname_lower.endswith(".docx")
        or fname_lower.endswith(".doc")
    )
    is_image = (
        mime_type.startswith("image/")
        or any(fname_lower.endswith(ext) for ext in (".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"))
    )

    try:
        if is_pdf:
            raw = _extract_pdf(file_bytes)
            # Determine which method was actually used
            # (_extract_pdf logs this; we re-check by char count for the return value)
            method = "pdfplumber"
            # If result looks like OCR output (very short lines, no spaces between words),
            # we report doctr_ocr — but simpler: just check if pdfplumber would have worked
            try:
                import pdfplumber
                test_parts = []
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            test_parts.append(t)
                if len("".join(test_parts).strip()) < MIN_TEXT_CHARS:
                    method = "doctr_ocr"
            except Exception:
                method = "doctr_ocr"

        elif is_docx:
            raw = _extract_docx(file_bytes)
            method = "python_docx"

        elif is_image:
            raw = _extract_image(file_bytes, mime_type)
            method = "doctr_image"

        else:
            raise ExtractionError(
                f"Unsupported file format: mime='{mime_type}', filename='{original_filename}'. "
                f"Supported: PDF, DOCX, JPG, PNG, TIFF."
            )

    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(
            f"Extraction failed for '{original_filename}': {exc}"
        ) from exc

    cleaned = _clean(raw)

    if not cleaned:
        raise ExtractionError(
            f"Extraction produced empty text for '{original_filename}'. "
            f"The file may be blank, password-protected, or corrupt."
        )

    logger.info(
        "Extracted %d chars from '%s' using %s.",
        len(cleaned), original_filename, method
    )
    return cleaned, method